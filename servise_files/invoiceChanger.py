import docx
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from allFinders import find_weight


def change_invoice(path):

    def set_cell_settings(cell_):
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].style.font.size = Pt(8)

    update_path = f'{os.path.split(path)[0]}\\Измененные накладные'
    new_invoice_path = f'{update_path}\\{os.path.basename(path)}'
    if os.path.exists(update_path):
        if os.path.isfile(new_invoice_path):
            print(f'Измененная накладная уже существует {new_invoice_path}')
            return
    else:
        os.mkdir(update_path)

    document = docx.Document(path)
    table = document.tables[1]
    max_row = len(table.rows)
    total_weight = 0

    for row in range(3, max_row - 21):
        string = table.rows[row].cells[4].text
        weight = find_weight(string.lower().split(' '))
        print(f'Поиск веса в строке ({string})')
        total_weight += float(weight)
        amount = (table.rows[row].cells[33].text.split(','))[0]

        cell = table.cell(row, 30)
        cell.text = amount
        set_cell_settings(cell)

        cell = table.cell(row, 33)
        cell.text = f'{weight} гр.'
        set_cell_settings(cell)

        cell = table.cell(row, 36)
        price = ''.join(((table.rows[row].cells[40].text.split(','))[0]).split('\xa0'))
        cell.text = str(round(float(price) / int(amount) / float(weight), 2)).replace('.', ',')
        while len(cell.text.split(',')[1]) < 2:
            cell.text += '0'
        set_cell_settings(cell)

    cell = table.cell(max_row - 21, 33)
    cell.text = f'{str(round(total_weight, 2))} гр.'
    set_cell_settings(cell)

    document.save(new_invoice_path)
    print(f'\nИзменения выполнены успешно. Измененная накладная сохранена в {new_invoice_path}\n')


# Test

# change_invoice('E:\Elena\Desktop\Документы  Александрова Е.П\Накладные\Накладные Каурова Евгения\
#                Расходная_накладная_№12_от_02-04-22.docx')
