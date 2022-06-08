import os
import docx
from handlerFunctions import find_art, find_description, find_weight
from validity import check_file_path


def docxParser(path):

    """ Функция извлечения данных из документа типа 'docx'. Построчно перебирает таблицу документа. Создает словари, где
     ключ, артикул изделия, а значения, описание изделия. Извлекает из строк описание изделия и добавляет их в
     словарь. Затем добавляет словари в список.
     Возвращает список словарей. """

    document = docx.Document(path)
    max_row = (len(document.tables[1].rows))
    group = 'word'
    art_list = []

    # Построчный проход по таблице накладной
    for i in range(3, max_row - 21):
        _string = document.tables[1].rows[i].cells[4].text
        split_string = _string.lower().split(' ')
        art = find_art(_string, group=group)
        art_dict = {find_art(_string, group=group): {}}
        art_dict[art]['Описание'] = find_description(split_string, group=group)
        art_dict[art]['Масса'] = str(find_weight(split_string, group=group)) + ' гр'

        art_list.append(art_dict)

    return art_list


def pdfParser(path, file):
    pass


def directory_parsing(path_to_invoices):

    """ Функция проходит по всем файлам и папкам находящимся в указанной директории, проверяет их на причастность к
    накладным и запускает для каждого подходящего файла парсеры соответствующие формату файла.  """

    file_list = []
    # Проход по файлам указанной директории
    for root, dirs, files in os.walk(path_to_invoices):
        for file in files:
            file_dict = {}

            if file.startswith('~'):
                continue

            path = path_to_invoices + '\\' + file

            if check_file_path(path) and file.endswith('.docx'):

                file_dict[file] = docxParser(path)
                file_list.append(file_dict)

                # if file.endswith('.pdf'):
                #     file_pdf_dictionary = pdfParser(path_to_invoices, file)
    for i in file_list:
        for key, values in i.items():
            print(key)
            counter = 0
            for value in values:
                counter += 1
                print(f'Строка {counter} === {value}')
    return file_list


# directory_parsing("E:\Elena\Desktop\Документы  Александрова Е.П\Накладные Кирпичникова Маргарита")
# directory_parsing('E:\Elena\Desktop\Документы  Александрова Е.П\Накладные\Накладные Каурова Евгения')
# directory_parsing("E:\Elena\Desktop\Документы  Александрова Е.П\Накладные Сартакова Светлана")

